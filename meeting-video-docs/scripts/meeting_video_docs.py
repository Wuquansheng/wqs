from __future__ import annotations

import argparse
import json
import re
from pathlib import Path


def require_deps():
    missing = []
    try:
        from faster_whisper import WhisperModel  # noqa: F401
    except ImportError:
        missing.append("faster-whisper")
    try:
        from docx import Document  # noqa: F401
    except ImportError:
        missing.append("python-docx")
    try:
        from opencc import OpenCC  # noqa: F401
    except ImportError:
        missing.append("opencc-python-reimplemented")
    if missing:
        raise SystemExit(
            "Missing dependencies: "
            + ", ".join(missing)
            + "\nInstall with: python -m pip install "
            + " ".join(missing)
        )


def parse_ts(value: str | None) -> float | None:
    if not value:
        return None
    parts = value.strip().split(":")
    if len(parts) == 3:
        h, m, s = parts
        return int(h) * 3600 + int(m) * 60 + float(s)
    if len(parts) == 2:
        m, s = parts
        return int(m) * 60 + float(s)
    return float(value)


def fmt_ts(seconds: float) -> str:
    seconds = max(0, int(round(seconds)))
    h = seconds // 3600
    m = (seconds % 3600) // 60
    s = seconds % 60
    return f"{h:02d}:{m:02d}:{s:02d}"


def transcribe(video: Path, out_dir: Path, model_name: str, language: str | None) -> Path:
    from faster_whisper import WhisperModel

    out_dir.mkdir(parents=True, exist_ok=True)
    segments_path = out_dir / f"{video.stem}_segments.json"
    model = WhisperModel(model_name, device="cpu", compute_type="int8")
    kwargs = {
        "beam_size": 5,
        "vad_filter": True,
        "vad_parameters": {"min_silence_duration_ms": 500},
    }
    if language:
        kwargs["language"] = language
    segments, info = model.transcribe(str(video), **kwargs)

    rows = []
    for seg in segments:
        text = seg.text.strip()
        if not text:
            continue
        rows.append({"start": seg.start, "end": seg.end, "text": text})
    metadata = {
        "video": str(video),
        "language": info.language,
        "language_probability": info.language_probability,
        "segments": rows,
    }
    segments_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")
    return segments_path


def load_segments(path: Path):
    from opencc import OpenCC

    data = json.loads(path.read_text(encoding="utf-8"))
    rows = data["segments"] if isinstance(data, dict) and "segments" in data else data
    cc = OpenCC("t2s")
    result = []
    for row in rows:
        text = cc.convert(row["text"].strip())
        if text:
            result.append({"start": row["start"], "end": row["end"], "text": text})
    return result


def set_font(run, size=10.5, bold=False):
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def set_defaults(doc):
    from docx.oxml.ns import qn
    from docx.shared import Pt

    for style_name, size, bold in [
        ("Normal", 10.5, False),
        ("Heading 1", 16, True),
        ("Heading 2", 14, True),
        ("Heading 3", 12, True),
    ]:
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.bold = bold


def add_title(doc, title: str, subtitle: str | None = None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_font(r, 18, True)
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        set_font(r)


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r)


def build_full_doc(segments, out_path: Path, title: str, qa_start: float | None):
    from docx import Document

    doc = Document()
    set_defaults(doc)
    add_title(doc, title, "完整版：简体中文，按演讲部分和问答部分整理")

    doc.add_heading("整理说明", level=1)
    add_bullet(doc, "本文件根据会议录屏机器转写整理，已转换为简体中文。")
    add_bullet(doc, "正文保留时间戳和原始口语表达；个别人名、软件名、专业名词可能存在识别误差。")
    if qa_start is None:
        add_bullet(doc, "未提供明确问答起点；以下内容按完整转写连续输出。")
        doc.add_heading("完整会议内容", level=1)
        selected = [("完整会议内容", segments)]
    else:
        add_bullet(doc, f"演讲/问答切分点：{fmt_ts(qa_start)}。")
        selected = [
            ("一、演讲部分：完整文字稿", [s for s in segments if s["start"] < qa_start]),
            ("二、问答部分：完整文字稿", [s for s in segments if s["start"] >= qa_start]),
        ]

    for idx, (heading, rows) in enumerate(selected):
        if idx:
            doc.add_page_break()
        doc.add_heading(heading, level=1)
        for seg in rows:
            p = doc.add_paragraph()
            r = p.add_run(f"[{fmt_ts(seg['start'])} - {fmt_ts(seg['end'])}] {seg['text']}")
            set_font(r)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def markdown_to_docx(md_path: Path, out_path: Path, title: str):
    from docx import Document

    doc = Document()
    set_defaults(doc)
    add_title(doc, title, "总结版：简体中文，分别归纳演讲部分和问答部分")
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif re.match(r"^[-*]\s+", line):
            add_bullet(doc, re.sub(r"^[-*]\s+", "", line))
        else:
            p = doc.add_paragraph()
            r = p.add_run(line)
            set_font(r)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def write_text_outputs(segments, out_dir: Path, stem: str):
    md = out_dir / f"{stem}_transcript_简体.md"
    txt = out_dir / f"{stem}_transcript_简体.txt"
    lines = [f"# {stem} 完整文字稿（简体）", ""]
    plain = []
    for seg in segments:
        lines.append(f"[{fmt_ts(seg['start'])} - {fmt_ts(seg['end'])}] {seg['text']}")
        plain.append(seg["text"])
    md.write_text("\n".join(lines) + "\n", encoding="utf-8")
    txt.write_text("\n".join(plain) + "\n", encoding="utf-8")


def main():
    parser = argparse.ArgumentParser(description="Create meeting transcript and Word documents from a recording.")
    parser.add_argument("--video", type=Path, help="Meeting video/audio file to transcribe.")
    parser.add_argument("--segments", type=Path, help="Existing segments JSON to reuse.")
    parser.add_argument("--out-dir", type=Path, required=True)
    parser.add_argument("--model", default="small")
    parser.add_argument("--language", default="zh")
    parser.add_argument("--qa-start", help="Actual Q&A start time, e.g. HH:MM:SS. Omit if unknown.")
    parser.add_argument("--summary-md", type=Path, help="Markdown summary prepared by Codex.")
    parser.add_argument("--title", default="会议纪要")
    args = parser.parse_args()

    require_deps()
    args.out_dir.mkdir(parents=True, exist_ok=True)
    if args.segments:
        segments_path = args.segments
    elif args.video:
        segments_path = transcribe(args.video, args.out_dir, args.model, args.language)
    else:
        raise SystemExit("Provide --video or --segments.")

    segments = load_segments(segments_path)
    stem = args.video.stem if args.video else segments_path.stem.replace("_segments", "")
    write_text_outputs(segments, args.out_dir, stem)
    qa_start = parse_ts(args.qa_start)
    build_full_doc(segments, args.out_dir / f"{stem}_完整版_简体.docx", args.title, qa_start)
    if args.summary_md:
        markdown_to_docx(args.summary_md, args.out_dir / f"{stem}_总结_简体.docx", args.title)
    print(f"segments={len(segments)}")
    print(f"full_docx={args.out_dir / f'{stem}_完整版_简体.docx'}")
    if args.summary_md:
        print(f"summary_docx={args.out_dir / f'{stem}_总结_简体.docx'}")


if __name__ == "__main__":
    main()
